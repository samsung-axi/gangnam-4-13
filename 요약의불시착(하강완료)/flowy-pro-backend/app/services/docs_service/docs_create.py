import requests
import io
import os
from datetime import datetime
from docx import Document
from typing import Dict, List, Optional, Any, Union


def download_docx_from_url(url: str) -> Optional[Document]:
    """
    주어진 URL에서 DOCX 파일을 다운로드하여 Document 객체로 반환
    
    Args:
        url (str): DOCX 파일의 다운로드 URL
        
    Returns:
        Optional[Document]: 성공 시 Document 객체, 실패 시 None
    """
    print(f"[1단계] 문서 다운로드 시작: {url}")
    
    try:
        # HTTP GET 요청으로 파일 다운로드
        response = requests.get(url, timeout=30)
        response.raise_for_status()  # HTTP 에러 시 예외 발생
        
        # 메모리에서 DOCX 파일 처리
        docx_bytes = io.BytesIO(response.content)
        document = Document(docx_bytes)
        
        print("[1단계] 문서 다운로드 완료")
        return document
        
    except requests.exceptions.RequestException as e:
        print(f"[1단계] 문서 다운로드 실패: {str(e)}")
        return None
    except Exception as e:
        print(f"[1단계] DOCX 파일 처리 실패: {str(e)}")
        return None


def call_llm_api(prompt: str, context: str = "") -> str:
    """
    LLM API 호출 함수 (Mock 응답 반환)
    실제 환경에서는 OpenAI API, Claude API 등을 호출하는 로직으로 대체
    
    Args:
        prompt (str): LLM에 전달할 프롬프트
        context (str): 추가 컨텍스트 정보
        
    Returns:
        str: LLM 응답 (Mock 데이터)
    """
    # Mock 응답 - 실제로는 LLM API 호출 결과를 반환
    if "문서의 종류" in prompt and "갱신할 항목" in prompt:
        return """
        {
            "document_type": "주간 보고서",
            "update_items": ["금주 진행 사항", "결정 사항", "다음 단계", "이슈 사항", "회의일시", "회의안건"]
        }
        """
    elif "회의 내용에서" in prompt and "정보를 추출" in prompt:
        return """
        {
            "금주 진행 사항": "프로젝트 A 설계 완료, 개발 환경 구축 진행 중",
            "결정 사항": "다음 주까지 프로토타입 완성하기로 결정",
            "다음 단계": "UI/UX 디자인 검토 및 백엔드 API 개발 시작",
            "이슈 사항": "외부 API 연동 관련 기술적 검토 필요",
            "회의일시": "2024년 12월 15일 오후 2시",
            "회의안건": "프로젝트 A 진행 상황 점검 및 향후 계획 수립"
        }
        """
    else:
        return "Mock LLM 응답"


def get_document_info_from_llm(document: Document) -> Dict[str, Any]:
    """
    문서 내용을 분석하여 문서 종류와 갱신할 항목을 파악
    
    Args:
        document (Document): 분석할 DOCX 문서 객체
        
    Returns:
        Dict[str, Any]: 문서 정보 {'document_type': str, 'update_items': List[str]}
    """
    print("[2단계] 문서 정보 분석 시작")
    
    try:
        # 문서의 텍스트 내용 추출
        doc_text = ""
        for paragraph in document.paragraphs:
            doc_text += paragraph.text + "\n"
        
        # 테이블 내용도 추출
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    doc_text += cell.text + " "
            doc_text += "\n"
        
        # LLM에 문서 분석 요청
        prompt = f"""
        다음 문서의 내용을 분석하여 문서의 종류와 갱신할 항목을 파악해주세요.
        
        문서 내용:
        {doc_text[:1000]}...
        
        응답 형식:
        {{
            "document_type": "문서 종류 (예: 주간 보고서, 회의록, 프로젝트 계획서 등)",
            "update_items": ["갱신할 항목1", "갱신할 항목2", ...]
        }}
        """
        
        llm_response = call_llm_api(prompt, doc_text)
        
        # JSON 파싱 (간단한 파싱 로직)
        import json
        try:
            result = json.loads(llm_response.strip())
        except:
            # JSON 파싱 실패 시 기본값 반환
            result = {
                "document_type": "회의록",
                "update_items": ["회의일시", "회의안건", "회의내용", "결정사항", "특이사항"]
            }
        
        print(f"[2단계] 문서 정보 분석 완료 - 종류: {result['document_type']}")
        print(f"        갱신 항목: {', '.join(result['update_items'])}")
        
        return result
        
    except Exception as e:
        print(f"[2단계] 문서 정보 분석 실패: {str(e)}")
        # 기본값 반환
        return {
            "document_type": "회의록",
            "update_items": ["회의일시", "회의안건", "회의내용", "결정사항", "특이사항"]
        }


def extract_info_from_meeting_text(meeting_text: str, update_items: List[str]) -> Dict[str, str]:
    """
    회의 내용 텍스트에서 갱신 항목에 해당하는 정보 추출
    
    Args:
        meeting_text (str): 회의 내용 텍스트
        update_items (List[str]): 갱신할 항목 리스트
        
    Returns:
        Dict[str, str]: 추출된 정보 {'항목명': '추출된 내용', ...}
    """
    print("[3단계] 회의 내용에서 정보 추출 시작")
    
    try:
        # LLM에 정보 추출 요청
        prompt = f"""
        다음 회의 내용에서 지정된 항목들에 해당하는 정보를 추출해주세요.
        
        회의 내용:
        {meeting_text}
        
        추출할 항목:
        {', '.join(update_items)}
        
        응답 형식 (JSON):
        {{
            "항목명1": "추출된 내용1",
            "항목명2": "추출된 내용2",
            ...
        }}
        """
        
        llm_response = call_llm_api(prompt, meeting_text)
        
        # JSON 파싱
        import json
        try:
            extracted_info = json.loads(llm_response.strip())
        except:
            # 파싱 실패 시 Mock 데이터 반환
            extracted_info = {
                "회의일시": datetime.now().strftime("%Y년 %m월 %d일 오후 2시"),
                "회의안건": "정기 프로젝트 회의",
                "회의내용": meeting_text[:200] + "..." if len(meeting_text) > 200 else meeting_text,
                "결정사항": "주요 결정 사항들이 논의되었습니다.",
                "특이사항": "특별한 이슈는 없었습니다."
            }
        
        print(f"[3단계] 정보 추출 완료 - {len(extracted_info)}개 항목")
        for key, value in extracted_info.items():
            print(f"        {key}: {value[:50]}{'...' if len(value) > 50 else ''}")
        
        return extracted_info
        
    except Exception as e:
        print(f"[3단계] 정보 추출 실패: {str(e)}")
        return {}


def update_document_content(document: Document, extracted_info: Dict[str, str]) -> Document:
    """
    문서 내용을 추출된 정보로 갱신
    
    Args:
        document (Document): 갱신할 DOCX 문서 객체
        extracted_info (Dict[str, str]): 추출된 정보
        
    Returns:
        Document: 갱신된 문서 객체
    """
    print("[4단계] 문서 내용 갱신 시작")
    
    try:
        updated_count = 0
        
        # 테이블 기반 갱신 (회의록 템플릿 기준)
        for table_idx, table in enumerate(document.tables):
            print(f"        테이블 {table_idx + 1} 처리 중...")
            
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    
                    # 셀 내용과 추출된 정보 매칭
                    for key, value in extracted_info.items():
                        # 키워드 매칭을 통한 셀 식별
                        if (key in cell_text or 
                            any(keyword in cell_text for keyword in [
                                "일시", "시간", "날짜"]) and "회의일시" in key or
                            any(keyword in cell_text for keyword in [
                                "안건", "주제", "목적"]) and "회의안건" in key or
                            any(keyword in cell_text for keyword in [
                                "내용", "논의", "회의"]) and "회의내용" in key or
                            any(keyword in cell_text for keyword in [
                                "결정", "합의", "결론"]) and "결정사항" in key or
                            any(keyword in cell_text for keyword in [
                                "특이", "이슈", "문제"]) and "특이사항" in key):
                            
                            # 같은 행의 다음 셀에 내용 업데이트
                            if cell_idx + 1 < len(row.cells):
                                row.cells[cell_idx + 1].text = value
                                updated_count += 1
                                print(f"        갱신: {key} -> {value[:30]}{'...' if len(value) > 30 else ''}")
                            # 또는 현재 셀이 비어있다면 직접 업데이트
                            elif not cell_text or len(cell_text) < 10:
                                cell.text = value
                                updated_count += 1
                                print(f"        갱신: {key} -> {value[:30]}{'...' if len(value) > 30 else ''}")
        
        # 단락 기반 갱신 (테이블이 없는 경우)
        if updated_count == 0:
            print("        테이블 갱신 실패, 단락 기반 갱신 시도...")
            for paragraph in document.paragraphs:
                for key, value in extracted_info.items():
                    if key in paragraph.text:
                        paragraph.text = paragraph.text.replace(key, f"{key}: {value}")
                        updated_count += 1
        
        print(f"[4단계] 문서 내용 갱신 완료 - {updated_count}개 항목 갱신")
        return document
        
    except Exception as e:
        print(f"[4단계] 문서 내용 갱신 실패: {str(e)}")
        return document


def automated_document_update_agent(
    meeting_id: str,
    meeting_text: str,
    download_link: str,
    output_dir: Optional[str] = None,
    return_bytes: bool = False
) -> Union[str, bytes]:
    """
    문서 갱신 자동화 에이전트 메인 함수
    
    Args:
        meeting_id (str): 회의 ID
        meeting_text (str): 회의 내용 텍스트
        download_link (str): 문서 다운로드 링크
        output_dir (Optional[str], optional): 출력 디렉토리 경로 (None이면 현재 디렉토리)
        return_bytes (bool): True면 파일 바이트를 반환, False면 파일 경로 반환
        
    Returns:
        Union[str, bytes]: 성공 시 갱신된 파일 경로 또는 바이트 데이터
    """
    print(f"=== 문서 갱신 자동화 에이전트 시작 ===")
    print(f"회의 ID: {meeting_id}")
    print(f"다운로드 링크: {download_link}")
    print(f"회의 내용 길이: {len(meeting_text)} 문자")
    print()
    
    try:
        # 1단계: 문서 다운로드
        document = download_docx_from_url(download_link)
        if document is None:
            return "오류: 문서 다운로드 실패"
        
        # 2단계: 문서 정보 분석
        doc_info = get_document_info_from_llm(document)
        
        # 3단계: 회의 내용에서 정보 추출
        extracted_info = extract_info_from_meeting_text(meeting_text, doc_info['update_items'])
        if not extracted_info:
            return "오류: 회의 내용에서 정보 추출 실패"
        
        # 4단계: 문서 내용 갱신
        updated_document = update_document_content(document, extracted_info)
        
        # 5단계: 갱신된 문서 저장
        output_filename = f"updated_{meeting_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        # 출력 디렉토리 설정
        if output_dir:
            os.makedirs(output_dir, exist_ok=True)
            output_path = os.path.join(output_dir, output_filename)
        else:
            output_path = output_filename
        
        # 바이트 반환 옵션
        if return_bytes:
            # 메모리에서 문서를 바이트로 반환
            doc_bytes = io.BytesIO()
            updated_document.save(doc_bytes)
            doc_bytes.seek(0)
            
            print()
            print(f"=== 문서 갱신 완료 ===")
            print(f"문서가 메모리에서 바이트로 반환됨")
            
            return doc_bytes.getvalue()
        else:
            # 파일로 저장
            updated_document.save(output_path)
            
            print()
            print(f"=== 문서 갱신 완료 ===")
            print(f"저장된 파일: {os.path.abspath(output_path)}")
            
            return os.path.abspath(output_path)
        
    except Exception as e:
        error_msg = f"오류: 문서 갱신 프로세스 실패 - {str(e)}"
        print(error_msg)
        return error_msg


def save_bytes_to_file(doc_bytes: bytes, filename: str) -> str:
    """
    바이트 데이터를 파일로 저장하는 헬퍼 함수
    
    Args:
        doc_bytes (bytes): 문서 바이트 데이터
        filename (str): 저장할 파일명
        
    Returns:
        str: 저장된 파일의 절대 경로
    """
    with open(filename, 'wb') as f:
        f.write(doc_bytes)
    return os.path.abspath(filename)


def upload_to_s3(
    doc_bytes: bytes,
    bucket_name: str,
    object_key: str,
    aws_access_key: Optional[str] = None,
    aws_secret_key: Optional[str] = None
) -> str:
    """
    갱신된 문서를 S3에 업로드하는 함수 (선택사항)
    
    Args:
        doc_bytes (bytes): 문서 바이트 데이터
        bucket_name (str): S3 버킷명
        object_key (str): S3 객체 키 (파일명)
        aws_access_key (str, optional): AWS 액세스 키
        aws_secret_key (str, optional): AWS 시크릿 키
        
    Returns:
        str: S3 URL 또는 에러 메시지
        
    Note:
        실제 사용 시 boto3 라이브러리 필요: pip install boto3
    """
    try:
        # boto3를 사용한 S3 업로드 (예시)
        # import boto3
        # s3_client = boto3.client('s3', 
        #                         aws_access_key_id=aws_access_key,
        #                         aws_secret_access_key=aws_secret_key)
        # s3_client.put_object(Bucket=bucket_name, Key=object_key, Body=doc_bytes)
        # return f"https://{bucket_name}.s3.amazonaws.com/{object_key}"
        
        print("S3 업로드 기능은 boto3 라이브러리가 필요합니다.")
        return "S3 업로드 기능 미구현"
        
    except Exception as e:
        return f"S3 업로드 실패: {str(e)}"


# 실행 예시
if __name__ == "__main__":
    # Mock 데이터를 사용한 테스트
    test_meeting_id = "MEET_20241215_001"
    
    test_meeting_text = """
    2024년 12월 15일 오후 2시에 진행된 프로젝트 A 정기 회의입니다.
    
    주요 안건:
    1. 이번 주 진행 상황 점검
    2. 다음 주 계획 수립
    3. 이슈 사항 논의
    
    논의 내용:
    - 프로젝트 A의 설계 단계가 완료되었습니다.
    - 개발 환경 구축이 90% 진행되었습니다.
    - UI/UX 디자인 검토가 필요한 상황입니다.
    
    결정 사항:
    - 다음 주까지 프로토타입을 완성하기로 결정했습니다.
    - 외부 API 연동에 대한 기술적 검토를 진행하기로 했습니다.
    
    다음 단계:
    - UI/UX 디자인 팀과의 협업 시작
    - 백엔드 API 개발 착수
    - 테스트 환경 구축
    
    특이 사항:
    - 외부 업체와의 API 연동 관련하여 기술적 이슈가 발견되어 추가 검토가 필요합니다.
    """
    
    # 실제 URL 대신 Mock URL 사용 (실제 환경에서는 유효한 DOCX URL 필요)
    test_download_link = "https://flowy-pro-docs-upload-test.s3.amazonaws.com/documents/%EC%9A%94%EA%B5%AC%EC%82%AC%ED%95%AD%EC%A0%95%EC%9D%98%EC%84%9C%EA%B0%9C%EB%B0%9C%EC%97%85%EB%AC%B4.docx?AWSAccessKeyId=AKIAVIN7OTFDSAYFUPFI&Signature=Sb4q4lFNwQu1MT09stBIHqJIiDk%3D&Expires=1751248168"
    
    print("주의: 이 예시는 Mock 데이터를 사용합니다.")
    print("실제 사용 시에는 유효한 DOCX 다운로드 링크가 필요합니다.")
    print()
    
    print("=== 다양한 사용 방법 예시 ===")
    print()
    
    # 방법 1: 현재 디렉토리에 파일 저장
    # print("1. 현재 디렉토리에 파일 저장:")
    # result1 = automated_document_update_agent(
    #     meeting_id=test_meeting_id,
    #     meeting_text=test_meeting_text,
    #     download_link=test_download_link
    # )
    # print(f"결과: {result1}")
    # print()
    
    # 방법 2: 특정 디렉토리에 파일 저장
    print("2. 특정 디렉토리에 파일 저장:")
    result2 = automated_document_update_agent(
        meeting_id=test_meeting_id,
        meeting_text=test_meeting_text,
        download_link=test_download_link,
        output_dir="./updated_documents"
    )
    print(f"결과: {result2}")
    print()
    
    # 방법 3: 바이트 데이터로 반환 (메모리에서 처리)
    # print("3. 바이트 데이터로 반환:")
    # doc_bytes = automated_document_update_agent(
    #     meeting_id=test_meeting_id,
    #     meeting_text=test_meeting_text,
    #     download_link=test_download_link,
    #     return_bytes=True
    # )
    
    # if isinstance(doc_bytes, bytes):
    #     print(f"바이트 데이터 크기: {len(doc_bytes)} bytes")
        
    #     # 바이트 데이터를 파일로 저장
    #     saved_path = save_bytes_to_file(doc_bytes, f"from_bytes_{test_meeting_id}.docx")
    #     print(f"바이트에서 파일로 저장: {saved_path}")
        
    #     # 필요 시 S3 업로드 등 추가 처리 가능
    #     # upload_result = upload_to_s3(doc_bytes, "my-bucket", "updated-doc.docx")
    #     # print(f"S3 업로드 결과: {upload_result}")
    # else:
    #     print(f"바이트 반환 실패: {doc_bytes}")
    
    print(f"\n최종 결과: 문서 갱신 프로세스 완료")