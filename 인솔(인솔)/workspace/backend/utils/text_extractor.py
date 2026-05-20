"""
다양한 파일 형식에서 텍스트를 추출하는 유틸리티
"""

import io
import logging
from typing import Optional, Tuple
from pathlib import Path

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False
    logging.warning("pdfplumber가 설치되지 않았습니다. PDF 처리가 제한됩니다.")

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx가 설치되지 않았습니다. DOCX 처리가 제한됩니다.")

try:
    from PyPDF2 import PdfReader
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False
    logging.warning("PyPDF2가 설치되지 않았습니다. PDF 처리가 제한됩니다.")

logger = logging.getLogger(__name__)


class TextExtractor:
    """다양한 파일 형식에서 텍스트를 추출하는 클래스"""
    
    def __init__(self):
        self.supported_formats = {
            '.txt': self._extract_text,
            '.pdf': self._extract_pdf,
            '.docx': self._extract_docx,
            '.doc': self._extract_doc,
        }
    
    def extract_text(self, file_bytes: bytes, filename: str) -> Tuple[str, str]:
        """
        파일에서 텍스트를 추출
        
        Args:
            file_bytes: 파일 바이트 데이터
            filename: 파일명
            
        Returns:
            (추출된 텍스트, 파일 타입)
        """
        try:
            file_path = Path(filename)
            file_extension = file_path.suffix.lower()
            
            if file_extension not in self.supported_formats:
                raise ValueError(f"지원하지 않는 파일 형식입니다: {file_extension}")
            
            extractor = self.supported_formats[file_extension]
            text = extractor(file_bytes, filename)
            
            # 텍스트 정제
            cleaned_text = self._clean_text(text)
            
            logger.info(f"텍스트 추출 완료: {filename} ({len(cleaned_text)} 문자)")
            return cleaned_text, file_extension
            
        except Exception as e:
            logger.error(f"텍스트 추출 실패: {filename}, 오류: {str(e)}")
            raise
    
    def _extract_text(self, file_bytes: bytes, filename: str) -> str:
        """일반 텍스트 파일 추출"""
        try:
            return file_bytes.decode('utf-8', errors='ignore')
        except UnicodeDecodeError:
            # UTF-8 실패 시 다른 인코딩 시도
            for encoding in ['cp949', 'euc-kr', 'latin-1']:
                try:
                    return file_bytes.decode(encoding, errors='ignore')
                except UnicodeDecodeError:
                    continue
            raise ValueError("텍스트 파일의 인코딩을 확인할 수 없습니다.")
    
    def _extract_pdf(self, file_bytes: bytes, filename: str) -> str:
        """PDF 파일에서 텍스트 추출"""
        if not PDFPLUMBER_AVAILABLE and not PYPDF2_AVAILABLE:
            raise ImportError("PDF 처리를 위한 라이브러리가 설치되지 않았습니다.")
        
        text_parts = []
        
        # pdfplumber 우선 사용 (더 정확한 텍스트 추출)
        if PDFPLUMBER_AVAILABLE:
            try:
                with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                    for page_num, page in enumerate(pdf.pages):
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(f"[페이지 {page_num + 1}]\n{page_text}")
                        else:
                            text_parts.append(f"[페이지 {page_num + 1}] - 텍스트 추출 불가")
                return "\n\n".join(text_parts)
            except Exception as e:
                logger.warning(f"pdfplumber로 PDF 추출 실패, PyPDF2로 재시도: {str(e)}")
        
        # PyPDF2로 대체
        if PYPDF2_AVAILABLE:
            try:
                pdf_reader = PdfReader(io.BytesIO(file_bytes))
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(f"[페이지 {page_num + 1}]\n{page_text}")
                    else:
                        text_parts.append(f"[페이지 {page_num + 1}] - 텍스트 추출 불가")
                return "\n\n".join(text_parts)
            except Exception as e:
                logger.error(f"PyPDF2로도 PDF 추출 실패: {str(e)}")
                raise
        
        raise ValueError("PDF 텍스트 추출에 실패했습니다.")
    
    def _extract_docx(self, file_bytes: bytes, filename: str) -> str:
        """DOCX 파일에서 텍스트 추출"""
        if not DOCX_AVAILABLE:
            raise ImportError("DOCX 처리를 위한 python-docx가 설치되지 않았습니다.")
        
        try:
            doc = docx.Document(io.BytesIO(file_bytes))
            text_parts = []
            
            # 제목 추출
            if doc.core_properties.title:
                text_parts.append(f"제목: {doc.core_properties.title}")
            
            # 단락 텍스트 추출
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # 표에서 텍스트 추출
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            return "\n".join(text_parts)
            
        except Exception as e:
            logger.error(f"DOCX 텍스트 추출 실패: {str(e)}")
            raise
    
    def _extract_doc(self, file_bytes: bytes, filename: str) -> str:
        """DOC 파일 처리 (현재는 지원하지 않음)"""
        raise ValueError("DOC 파일 형식은 현재 지원하지 않습니다. DOCX로 변환 후 업로드해주세요.")
    
    def _clean_text(self, text: str) -> str:
        """추출된 텍스트 정제"""
        if not text:
            return ""
        
        # 불필요한 공백 정리
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            # 빈 줄 제거
            if line.strip():
                # 앞뒤 공백 제거
                cleaned_line = line.strip()
                # 연속된 공백을 하나로
                cleaned_line = ' '.join(cleaned_line.split())
                cleaned_lines.append(cleaned_line)
        
        return '\n'.join(cleaned_lines)
    
    def get_file_info(self, file_bytes: bytes, filename: str) -> dict:
        """파일 정보 반환"""
        file_path = Path(filename)
        file_size = len(file_bytes)
        file_extension = file_path.suffix.lower()
        
        return {
            "filename": filename,
            "file_size": file_size,
            "file_extension": file_extension,
            "is_supported": file_extension in self.supported_formats
        }
    
    def validate_file(self, file_bytes: bytes, filename: str) -> bool:
        """파일 유효성 검사"""
        try:
            file_info = self.get_file_info(file_bytes, filename)
            
            if not file_info["is_supported"]:
                return False
            
            # 파일 크기 제한 (50MB)
            if file_info["file_size"] > 50 * 1024 * 1024:
                return False
            
            # 최소 파일 크기 (1KB)
            if file_info["file_size"] < 1024:
                return False
            
            return True
            
        except Exception:
            return False


# 전역 인스턴스
text_extractor = TextExtractor()


def extract_text_from_file(file_bytes: bytes, filename: str) -> Tuple[str, str]:
    """파일에서 텍스트 추출 (편의 함수)"""
    return text_extractor.extract_text(file_bytes, filename)


def validate_upload_file(file_bytes: bytes, filename: str) -> bool:
    """업로드 파일 유효성 검사 (편의 함수)"""
    return text_extractor.validate_file(file_bytes, filename)
