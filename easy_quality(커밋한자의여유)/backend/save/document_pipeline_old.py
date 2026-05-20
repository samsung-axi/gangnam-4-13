"""
LangGraph 기반 문서 처리 파이프라인 v9.1

v9.1 개선:
- 페이지 번호 메타데이터 추가
- Parent-Child 계층 구조 도입
- 문서 헤더 메타데이터 추출 (SOP ID, Version, Date, Department)
- 목차 중복 제거
- 메타데이터 최적화 (불필요 필드 제거)

노드 흐름:
┌─────────┐    ┌─────────┐    ┌──────────┐    ┌─────────┐
│  Load   │───▶│ Convert │───▶│ Validate │───▶│  Split  │
└─────────┘    └─────────┘    └──────────┘    └─────────┘
                    │              │                │
                    ▼              ▼                ▼
              ┌──────────┐  ┌──────────┐    ┌──────────┐
              │ Fallback │  │  Repair  │    │ Optimize │
              └──────────┘  └──────────┘    └──────────┘
                                                   │
                                                   ▼
                                            ┌──────────┐
                                            │ Finalize │
                                            └──────────┘
"""

from typing import TypedDict, List, Dict, Optional, Literal, Annotated
from dataclasses import dataclass, field
import re
from io import BytesIO
import operator
import json
from concurrent.futures import ThreadPoolExecutor, as_completed
from ..llm import get_llm_response


# ═══════════════════════════════════════════════════════════════════════════
# 상태 정의
# ═══════════════════════════════════════════════════════════════════════════

class PipelineState(TypedDict):
    """파이프라인 상태"""
    # 입력
    filename: str
    content: bytes

    # 설정
    chunk_size: int
    chunk_overlap: int
    use_llm_metadata: bool  # LLM 메타데이터 추출 사용 여부
    use_clause_parsing: bool  # 조항 번호 기반 파싱 사용 여부

    # 중간 결과
    file_type: str
    markdown: str
    metadata: Dict
    sections: List[Dict]
    chunks: List[Dict]

    # 품질 지표
    quality_score: float
    conversion_method: str

    # 에러 처리
    errors: Annotated[List[str], operator.add]
    warnings: Annotated[List[str], operator.add]
    retry_count: int

    # 최종 결과
    success: bool


@dataclass
class Chunk:
    """청크 데이터"""
    text: str
    index: int = 0
    metadata: Dict = field(default_factory=dict)


# ═══════════════════════════════════════════════════════════════════════════
# 문서 메타데이터 추출
# ═══════════════════════════════════════════════════════════════════════════

def extract_document_metadata(text: str, filename: str) -> Dict:
    """
    문서 헤더에서 메타데이터 추출 (Pure AI: LLM Only)
    """
    print(f"🧠 [Metadata] AI 기반 지능형 메타데이터 추출 중... (File: {filename})")
    
    head_text = text[:3000]
    
    prompt = f"""당신은 GMP 규정(SOP) 분석 전문가입니다. 다음 문서의 내용을 분석하여 관리용 메타데이터를 추출하세요.

[추출 규칙]
1. doc_id: 'EQ-SOP-0001'과 같은 관리 번호를 찾으세요.
2. title: 문서의 공식 명칭을 정확히 추출하세요. 
3. version: '1.0' 또는 'Ver 2.1' 같은 형식을 찾으세요.
4. effective_date: YYYY-MM-DD 형식으로 변환하세요.
5. owning_dept: 생산팀, 품질보증팀 등 담당 조직명을 찾으세요.

결과는 반드시 아래 JSON 형식으로만 답변하세요.
{{
  "doc_id": "추출된 ID (없으면 null)",
  "title": "문서 제목 (필수)",
  "version": "추출된 버전 (없으면 null)",
  "effective_date": "YYYY-MM-DD (없으면 null)",
  "owning_dept": "담당 부서 (없으면 null)"
}}

[파일명]
{filename}

[주의사항]
반드시 생각 과정(Reasoning)을 생략하고, 최종 JSON 데이터만 출력하세요.

[문서 내용]
{head_text}"""
    
    metadata = {"file_name": filename}
    try:
        # GPT-4o-mini: JSON 추출에 최적화
        llm_res = get_llm_response(prompt, llm_model="gpt-4o-mini", max_tokens=4096, temperature=0.1)
        json_match = re.search(r'\{.*\}', llm_res, re.DOTALL)
        if json_match:
            llm_meta = json.loads(json_match.group(0))
            # 호환성 레이어: doc_id를 sop_id로도 복사
            if 'doc_id' in llm_meta and 'sop_id' not in llm_meta:
                llm_meta['sop_id'] = llm_meta['doc_id']
            metadata.update(llm_meta)
            print(f"✅ [Metadata] AI 추출 성공: {metadata.get('doc_id') or metadata.get('sop_id') or 'ID 미확인'}")
    except Exception as e:
        print(f"⚠️ [Metadata] AI 추출 실패: {e}")
        metadata.update({"doc_id": None, "sop_id": None, "title": filename, "version": None, "effective_date": None, "owning_dept": None})

    return metadata

def extract_clause_metadata(text: str, doc_info: Dict, section_name: str) -> Dict:
    """
    조항(Clause) 단위 상세 메타데이터 추출 (processor 방식)
    """
    # 너무 짧은 내용은 분석 스킵
    clean_text = text.strip()
    if len(clean_text) < 30:
        return {}

    # print(f"🧠 [Clause Scan] {section_name} 상세 분석 중...")

    prompt = f"""당신은 GMP 문서 분석 전문가입니다.
다음 문서 청크를 분석해서 검색에 유용한 메타데이터를 추출하세요.

## 청크 정보
- 조항번호: {section_name}
- 제목: {section_name}
- 내용:
{text[:1000]}

## 추출할 메타데이터
다음 항목들을 JSON 형식으로 추출하세요:

1. content_type: 이 청크가 설명하는 내용의 유형 (예: 목적, 정의, 책임, 절차, 기준, 기록, 참고문헌 등)
2. main_topic: 이 청크의 핵심 주제
3. sub_topics: 관련 세부 주제들 (리스트, 최대 3개)
4. actors: 언급된 역할자/담당자 (리스트)
5. actions: 수행해야 하는 행위/절차 (리스트, 최대 3개)
6. conditions: 특수 조건이나 상황 (리스트)
7. summary: 한 문장 요약 (30자 이내)
8. intent_scope: 이 조항이 다루는 관리 영역 (다음 중 1개만 선택)
   - user_account: 사용자 계정·권한·역할 관리
   - document_lifecycle: 문서의 수명주기 (작성, 승인, 개정, 폐기 등)
   - system_configuration: 시스템 설정 or 구조 변경
   - audit_evidence: 감사대응에 관련한 자료
   - training: 교육, 훈련, 자격
9. intent_summary: 이 조항이 어떤 질문에 답하는지를 영어 1문장으로 요약 (예: "What is the purpose of Level 1 quality manual?")
10. language: 이 청크의 주요 언어 ("ko" 또는 "en")

## 출력
JSON만 출력:
{{"content_type": "...", "main_topic": "...", "sub_topics": [...], "actors": [...], "actions": [...], "conditions": [...], "summary": "...", "intent_scope": "...", "intent_summary": "...", "language": "..."}}"""

    try:
        # GPT-4o-mini: 빠르고 정확한 조항별 메타데이터 JSON 추출
        llm_res = get_llm_response(prompt, llm_model="gpt-4o-mini", max_tokens=4096, temperature=0.1)
        # processor 방식: JSON 파싱
        result_text = llm_res.strip()

        # JSON 코드 블록 처리
        if "```" in result_text:
            json_match = result_text.split("```")[1]
            if json_match.startswith("json"):
                json_match = json_match[4:]
            result_text = json_match.strip()

        res = json.loads(result_text)

        # 리스트를 쉼표 구분 문자열로 변환 (기존 pipeline 방식과 호환)
        if isinstance(res.get('sub_topics'), list):
            res['sub_topics'] = ', '.join(res['sub_topics'])
        if isinstance(res.get('actors'), list):
            res['actors'] = ', '.join(res['actors'])
        if isinstance(res.get('actions'), list):
            res['actions'] = ', '.join(res['actions'])
        if isinstance(res.get('conditions'), list):
            res['conditions'] = ', '.join(res['conditions'])

        # 호환성 보장
        if 'doc_id' not in res and doc_info.get('doc_id'):
            res['doc_id'] = doc_info.get('doc_id')

        return res
    except json.JSONDecodeError as e:
        print(f"⚠️ [Clause Scan] JSON 파싱 실패: {e}")
        return {
            "content_type": "",
            "main_topic": "",
            "sub_topics": "",
            "actors": "",
            "actions": "",
            "conditions": "",
            "summary": "",
            "intent_scope": "",
            "intent_summary": "",
            "language": "ko"
        }
    except Exception as e:
        print(f"⚠️ [Clause Scan] 실패: {e}")
        return {}

    return {}


# ═══════════════════════════════════════════════════════════════════════════
# 노드 함수들
# ═══════════════════════════════════════════════════════════════════════════

def node_load(state: PipelineState) -> PipelineState:
    """
    1단계: 파일 로드 및 타입 감지
    """
    filename = state["filename"]
    content = state["content"]
    
    # 실제 확장자 추출 (마지막 . 이후)
    filename_lower = filename.lower()
    if '.' in filename_lower:
        actual_ext = filename_lower.rsplit('.', 1)[-1]
    else:
        actual_ext = ''
    
    # 확장자 기반 타입 결정
    if actual_ext == 'pdf':
        file_type = 'pdf'
    elif actual_ext == 'docx':
        file_type = 'docx'
    elif actual_ext == 'doc':
        file_type = 'doc'
    elif actual_ext in ['html', 'htm']:
        file_type = 'html'
    elif actual_ext == 'md':
        file_type = 'markdown'
    elif actual_ext == 'txt':
        file_type = 'text'
    else:
        file_type = 'unknown'
    
    # 파일 크기 확인
    file_size = len(content)
    if file_size == 0:
        state["errors"] = ["파일이 비어있습니다."]
        state["success"] = False
        return state
    
    state["file_type"] = file_type
    state["metadata"] = {
        "file_name": filename,
        "file_type": file_type,
        "file_size": file_size,
    }
    
    return state


def node_convert(state: PipelineState) -> PipelineState:
    """
    2단계: 문서 → 마크다운 변환
    """
    file_type = state["file_type"]
    content = state["content"]
    filename = state["filename"]
    
    try:
        if file_type == 'docx':
            markdown, metadata = _convert_docx(filename, content)
            state["conversion_method"] = "python-docx"
            
        elif file_type == 'pdf':
            markdown, metadata, method = _convert_pdf_with_fallback(filename, content)
            state["conversion_method"] = method
            # V22.9 Global Noise Filter: 헤더 추론 및 분할 전 전체 문서 클리닝
            markdown = _clean_noise_globally(markdown)
            # PDF는 헤더 추론 필수!
            markdown = _infer_headers(markdown)
            state["conversion_method"] += "+infer-headers"
            
        elif file_type == 'html':
            markdown, metadata = _convert_html(filename, content)
            state["conversion_method"] = "beautifulsoup"
            
        elif file_type == 'markdown':
            markdown = content.decode('utf-8', errors='ignore')
            metadata = {}
            state["conversion_method"] = "passthrough"
            
        elif file_type == 'text':
            markdown = _convert_text_to_markdown(content.decode('utf-8', errors='ignore'))
            metadata = {}
            state["conversion_method"] = "text-inference"
            
        else:
            markdown = content.decode('utf-8', errors='ignore')
            metadata = {}
            state["conversion_method"] = "fallback-text"
            state["warnings"] = [f"알 수 없는 파일 타입: {file_type}, 텍스트로 처리"]
        
        # 문서 메타데이터 추출 (옵션)
        if state.get("use_llm_metadata", False):
            doc_meta = extract_document_metadata(markdown, filename)
            metadata.update(doc_meta)
        else:
            print(f"   ℹ️ LLM 메타데이터 추출 비활성화 (기본 메타데이터만 사용)")
        
        state["markdown"] = markdown
        state["metadata"].update(metadata)
        
    except Exception as e:
        state["errors"] = [f"변환 실패: {str(e)}"]
        state["markdown"] = ""
    
    return state


def node_convert_fallback(state: PipelineState) -> PipelineState:
    """
    2-1단계: 변환 실패 시 폴백 전략
    """
    content = state["content"]
    file_type = state["file_type"]
    
    state["warnings"] = [f"기본 변환 실패, 폴백 전략 시도 중..."]
    
    try:
        if file_type == 'pdf':
            markdown = _pdf_fallback_extract(content)
            state["conversion_method"] = "pdf-fallback"
            
        elif file_type == 'docx':
            markdown = _docx_fallback_extract(content)
            state["conversion_method"] = "docx-fallback"
            
        else:
            markdown = content.decode('utf-8', errors='ignore')
            state["conversion_method"] = "binary-text"
        
        state["markdown"] = markdown
        state["errors"] = []
        
    except Exception as e:
        state["errors"] = [f"폴백 변환도 실패: {str(e)}"]
        state["success"] = False
    
    return state


def node_validate(state: PipelineState) -> PipelineState:
    """
    3단계: 마크다운 품질 검증
    """
    markdown = state.get("markdown", "")
    
    # [최적화] 변환이 완료되었으므로 대용량 바이너리 데이터는 삭제 (UI 속도 향상)
    if markdown and "content" in state:
        state["content"] = b"" # 메모리 및 UI 렌더링 부하 감소
    
    if not markdown:
        state["quality_score"] = 0.0
        state["errors"] = ["마크다운 변환 결과가 비어있습니다."]
        return state
    
    score = 0.0
    issues = []
    
    # 1. 길이 체크 (최소 100자)
    if len(markdown) >= 100:
        score += 0.2
    else:
        issues.append("텍스트가 너무 짧음")
    
    # 2. 헤더 존재 여부
    header_count = len(re.findall(r'^#{1,6}\s+', markdown, re.MULTILINE))
    if header_count >= 3:
        score += 0.3
    elif header_count >= 1:
        score += 0.15
        issues.append("헤더가 부족함")
    else:
        issues.append("헤더가 없음")
    
    # 3. 문단 구조
    paragraphs = [p for p in markdown.split('\n\n') if p.strip()]
    if len(paragraphs) >= 5:
        score += 0.2
    elif len(paragraphs) >= 2:
        score += 0.1
        issues.append("문단 구조가 부실함")
    
    # 4. 한글 비율
    korean_chars = len(re.findall(r'[가-힣]', markdown))
    total_chars = len(markdown)
    korean_ratio = korean_chars / total_chars if total_chars > 0 else 0
    
    if korean_ratio >= 0.1:
        score += 0.2
    else:
        issues.append("한글 비율이 낮음")
    
    # 5. 특수문자 오염 체크
    garbage_ratio = len(re.findall(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', markdown)) / len(markdown) if markdown else 0
    if garbage_ratio < 0.01:
        score += 0.1
    else:
        issues.append("특수문자 오염 감지")
    
    state["quality_score"] = min(score, 1.0)
    
    if issues:
        state["warnings"] = issues
    
    return state


def node_repair(state: PipelineState) -> PipelineState:
    """
    3-1단계: 마크다운 품질 보정
    """
    markdown = state.get("markdown", "")
    
    state["warnings"] = ["품질 보정 수행 중..."]
    state["retry_count"] = state.get("retry_count", 0) + 1
    
    # 1. 특수문자 제거
    markdown = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', markdown)
    
    # 2. 연속 빈 줄 정리
    markdown = re.sub(r'\n{3,}', '\n\n', markdown)
    
    # 3. 헤더가 없으면 추론해서 추가
    if not re.search(r'^#{1,6}\s+', markdown, re.MULTILINE):
        markdown = _infer_headers(markdown)
    
    # 4. 깨진 테이블 복구 시도
    markdown = _repair_tables(markdown)
    
    state["markdown"] = markdown
    state["conversion_method"] += "+repaired"
    
    # 품질 재측정
    state = node_validate(state)
    
    return state


# ═══════════════════════════════════════════════════════════════════════════
# 조항 추출 엔진 (Pattern Registry)
# ═══════════════════════════════════════════════════════════════════════════

CLAUSE_PATTERNS = [
    # 1. 숫자 계층형 (1., 1.1, 5.2.1)
    {
        "name": "numeric_dot",
        "pattern": re.compile(r'^(\d+(?:\.\d+)*)\.?\s+(.*?)$'), # 마침표 뒤 공백 필수 (1. 제목)
        "level_func": lambda m: m.count('.')
    },
    # 1.1 숫자 단순형 (마침표 없는 번호 + 제목) -> 1 제목
    {
        "name": "numeric_simple",
        "pattern": re.compile(r'^(\d+)\s+([가-힣\w].*)$'), 
        "level_func": lambda m: 0
    },
    # 2. 한국어 법령/규정형 (제1조, 제12조)
    {
        "name": "ko_article",
        "pattern": re.compile(r'^(제\s*\d+\s*조)\s*(.*?)$'),
        "level_func": lambda m: 0
    },
    # 3. 한국어 순서형 (가., 나., 다.)
    {
        "name": "ko_alphabet",
        "pattern": re.compile(r'^([가-힣])\.\s*(.*?)$'),
        "level_func": lambda m: 1
    },
    # 4. 괄호 숫자형 ((1), (2) 또는 1), 2))
    {
        "name": "bracket_numeric",
        "pattern": re.compile(r'^\(?(\d+)\)\s*(.*?)$'),
        "level_func": lambda m: 2
    },
    # 5. 원문자 숫자형 (①, ②)
    {
        "name": "circle_numeric",
        "pattern": re.compile(r'^([①-⑳])\s*(.*?)$'),
        "level_func": lambda m: 2
    },
    # 6. 영문 대문자형 (A., B., C.)
    {
        "name": "en_uppercase",
        "pattern": re.compile(r'^([A-Z])\.\s*(.*?)$'),
        "level_func": lambda m: 1
    }
]

def split_by_clause(markdown: str, max_level: int = 0) -> List[Dict]:
    """
    유연한 조항 추출 엔진: 정규식 레지스트리를 순회하며 매칭되는 모든 조항 분할
    """
    lines = markdown.split('\n')
    clauses = []
    
    # 현재 매칭된 패턴 이름 (문서 내 일관성 유지 시도)
    active_pattern_names = set()

    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue

        # 헤더 마커 제거
        clean_line = re.sub(r'^#+\s*', '', line)
        
        # 패턴 매칭 시도
        match = None
        matched_pattern = None
        for p in CLAUSE_PATTERNS:
            m = p["pattern"].match(clean_line)
            if m:
                # 조항 번호 뒤에 제목이 없더라도 인식 (번호만 있는 라인 방어)
                num = m.group(1)
                title = m.group(2).strip() or f"Section {num}"
                match = m
                matched_pattern = p
                break
        
        if match:
            clause_num = match.group(1)
            title = match.group(2).strip() or f"Section {clause_num}"
            level = matched_pattern["level_func"](clause_num)
            
            # max_level 필터 (0이면 무제한)
            if max_level > 0 and level > max_level:
                i += 1
                continue
                
            content_lines = []
            i += 1
            
            # 다음 조항이 나올 때까지 수집
            while i < len(lines):
                next_line = lines[i].strip()
                if not next_line:
                    content_lines.append(lines[i])
                    i += 1
                    continue
                    
                next_clean = re.sub(r'^#+\s*', '', next_line)
                
                # 다음 라인이 어떤 조항 패턴에 걸리는지 확인
                is_next_clause = False
                for p in CLAUSE_PATTERNS:
                    if p["pattern"].match(next_clean):
                        is_next_clause = True
                        break
                
                if is_next_clause:
                    break
                    
                content_lines.append(lines[i])
                i += 1
            
            content = '\n'.join(content_lines).strip()
            
            # 계층 구조 (숫자형일 때만 부모 계산)
            parent_clauses = []
            if matched_pattern["name"] == "numeric_dot":
                parts = clause_num.split('.')
                for j in range(len(parts) - 1):
                    parent_clauses.append('.'.join(parts[:j+1]))
            
            clauses.append({
                "clause": clause_num,
                "title": title,
                "content": content,
                "level": level,
                "parent_clauses": parent_clauses,
                "pattern": matched_pattern["name"]
            })
        else:
            i += 1
            
    return clauses


def detect_content_start_with_llm(markdown: str) -> Optional[str]:
    """
    LLM을 사용하여 문서 헤더/목차 구간이 끝나고 실제 본문 조항이 시작되는 지점의 텍스트(Anchor)를 찾습니다.
    (예: "5.10.3", "제1조" 등)
    """
    sample_text = markdown[:5000]
    
    prompt = f"""당신은 GMP 문서 구조 분석 전문가입니다.
다음 마크다운 문서 내용을 보고, 문서 정보(SOP Number, Version 등)와 목차(ToC)가 완전히 종료되고,
실제 실행 지침이나 조항(Clause)이 본격적으로 시작되는 첫 번째 문장이나 조항 번호를 찾으세요.

[작업 규칙]
1. 본문이 시작되는 지점의 텍스트(예: "5.1", "제1조", "1. 목적")를 정확히 한 문장만 출력하세요.
2. 만약 문서의 처음부터 바로 본문이 시작된다면 "START"라고 답변하세요.
3. 생각 과정 없이 결과 텍스트만 출력하세요.

[문서 샘플]
{sample_text}
"""
    try:
        # Z.AI: 한국어 문서 구조 및 조항 시작 지점 파악에 탁월
        anchor = get_llm_response(prompt, llm_backend="zai", max_tokens=100, temperature=0.1).strip()
        if not anchor or "START" in anchor.upper():
            return None
        # Anchor가 너무 길면 무시 (정확한 매칭을 위해)
        if len(anchor) > 100:
            return None
        return anchor
    except:
        return None

def discover_structure_with_llm(markdown: str) -> List[Dict]:
    """
    AI를 사용하여 문서 전체의 조항 지도를 생성합니다.
    """
    sample_text = markdown[:40000] 
    
    prompt = f"""당신은 전문 문서 구조화 엔진입니다.
다음 마크다운 문서에서 '모든 조항'의 시작 지점을 찾아 JSON 리스트로 응답하세요.

[출력 형식]
[
  {{
    "clause": "번호 (예: 5.1.2.1)", 
    "title": "제목 (제목이 없다면 번호를 다시 적으세요. 문장은 넣지 마세요.)", 
    "level": 계층레벨(0부터 시작),
    "anchor": "해당 조항이 시작되는 텍스트 (반드시 조항 번호를 포함해야 함, 예: '5.1.2 본 절차서는')"
  }},
  ...
]

[작업 규칙]
1. 문서에 존재하는 모든 번호가 매겨진 조항(2.1, 2.2 등 아주 세밀한 단위까지)을 빠짐없이 포함하세요.
2. 'anchor'는 본문에 존재하는 텍스트와 완벽히 일치해야 하며, 반드시 '조항 번호'로 시작해야 합니다.
3. 2.1 뒤에 바로 2.2가 나오더라도 각각 별도의 앵커로 잡으세요. (가장 중요)
4. 'title'은 5단어 이하의 명칭이어야 합니다. 문장(예: ~적용된다, ~한다)은 절대 제목으로 추출하지 마세요.
5. 'page X of Y', 'Number: SOP-...' 같은 페이지 헤더/푸터 정보는 무시하세요.
6. JSON 외의 텍스트는 출력하지 마세요.

[문서 내용]
{sample_text}
"""
    try:
        llm_res = get_llm_response(prompt, llm_backend="zai", max_tokens=4000, temperature=0.1)
        json_match = re.search(r'\[.*\]', llm_res, re.DOTALL)
        if json_match:
            return json.loads(json_match.group(0))
    except Exception as e:
        print(f"   ⚠️ [AI Map] AI 지도 생성 실패: {e}")
    return []

def node_split(state: PipelineState) -> PipelineState:
    """
    3단계: AI 주도 하이브리드 조항 분할 (AI-Led Hybrid Parsing)
    1. AI가 먼저 문서의 전체 구조(ID, Title)와 시작점(Anchor)을 파악합니다.
    2. 파악된 앵커를 기준으로 본문을 정밀하게 분할합니다.
    """
    markdown = state.get("markdown", "")
    use_clause_parsing = state.get("use_clause_parsing", True)
    
    if not markdown:
        state["sections"] = []
        return state

    # 1. 지능형 시작 지점 식별 (이미 수행됨)
    effective_markdown = markdown
    start_anchor = detect_content_start_with_llm(markdown)
    if start_anchor:
        anchor_idx = markdown.find(start_anchor)
        if anchor_idx >= 0:
            effective_markdown = markdown[anchor_idx:]

    if use_clause_parsing:
        print("   [AI Map] AI를 사용하여 조항 지도를 생성 중...")
        # AI가 본문 전체를 훑으며 조항의 위치(앵커)를 찾습니다.
        structure = discover_structure_with_llm(effective_markdown)
        
        if structure:
            sections = []
            # 앵커를 기준으로 본문을 자릅니다.
            for i, item in enumerate(structure):
                anchor = item.get("anchor", "")
                next_anchor = structure[i+1].get("anchor", "") if i+1 < len(structure) else None
                
                # 앵커 위치 찾기
                start_pos = effective_markdown.find(anchor) if anchor else -1
                if start_pos == -1: continue # 앵커를 못 찾으면 스킵 (AI가 잘못 생성한 경우)
                
                # 다음 앵커 전까지를 내용으로 간주
                if next_anchor:
                    end_pos = effective_markdown.find(next_anchor, start_pos + len(anchor))
                    if end_pos == -1: end_pos = len(effective_markdown)
                else:
                    end_pos = len(effective_markdown)
                
                content = effective_markdown[start_pos:end_pos].strip()
                
                # [V22.6 Content Purity] 반복적인 SOP 헤더/푸터 및 메타데이터 제거
                junk_patterns = [
                    r'Number:\s*EQ-SOP-\d+.*',
                    r'Version:\s*\d+\.\d+.*',
                    r'Effective Date:\s*\d{4}-\d{2}-\d{2}.*',
                    r'Owning Department\s*.*',
                    r'Title\s*.*',
                    r'\d+\s+of\s+\d+', # 페이지 번호 (9 of 13)
                    r'GMP\s+문서\s+체계.*' # 문서명 반복 노출
                ]
                for pattern in junk_patterns:
                    content = re.sub(pattern, '', content, flags=re.IGNORECASE).strip()

                # 섹션 객체 생성
                clause_num = item.get("clause", "")
                title = item.get("title", "Untitled").strip()
                level = item.get("level", 0)
                
                # [V22.7 Title Refinement] 문장형 제목 감지 및 정규화 (Heuristics)
                # 1. 문장 종결 어미 체크 (다., 함., 임. 등)
                is_sentence = any(title.endswith(x) for x in ["다.", "함.", "임.", "요.", "다", "함", "임"])
                # 2. 한국어 조사 포함 여부 (는, 은, 이, 가, 를, 을 - 주어나 목적어가 있는 문장일 확률 높음)
                has_particles = any(x in title for x in ["는 ", "은 ", "이 ", "가 ", "를 ", "을 "])
                
                # 제목이 문장형이거나 길이가 30자를 넘는 경우 정규화
                if clause_num and (len(title) > 30 or is_sentence or (has_particles and len(title) > 20)):
                    print(f"   🩹 [Title] 문장형 제목 정규화: {title[:20]}...")
                    # 원래 제목이 본문에 아직 포함되지 않은 경우에만 본문 상단에 추가 (중복 방지)
                    if title not in content:
                        content = f"{title}\n\n{content}"
                    title = f"조항 {clause_num}"
                
                # [Junk Merging] 조항 번호가 없고 제목이 너무 길면(50자 이상) 본문 파편으로 간주하여 병합
                if not clause_num and len(title) > 50:
                    if sections:
                        # 이전 섹션에 내용 추가
                        sections[-1]["content"] += f"\n\n{content}"
                        print(f"   🩹 [Merge] 파편섹션('{title[:20]}...')을 이전 조항에 병합했습니다.")
                        continue
                
                # [V22.8 Sub-split] AI가 놓친 세부 조항(2.2 등)이 본문 중간에 섞여있는지 정규식으로 한 번 더 체크
                sub_sections = []
                if clause_num and "." in clause_num:
                    # 현재 조항 번호 형식을 기반으로 비슷한 패턴 탐색 (예: 2.1 -> 2.2, 2.3 등)
                    base_num = ".".join(clause_num.split(".")[:-1])
                    # 2.x 패턴 탐색
                    pattern = rf'\n({re.escape(base_num)}\.\d+)\s+'
                    parts = re.split(pattern, "\n" + content)
                    
                    if len(parts) > 1:
                        # 첫 번째 파트는 원래 조항의 내용
                        content = parts[0].strip()
                        # 이후 파트들은 번호, 내용, 번호, 내용... 순서임
                        for j in range(1, len(parts), 2):
                            s_num = parts[j]
                            s_content = parts[j+1].strip() if j+1 < len(parts) else ""
                            sub_sections.append({
                                "clause": s_num,
                                "title": f"조항 {s_num}",
                                "content": s_content,
                                "level": level,
                                "is_sub_split": True
                            })

                # [Normalization] 헤더 구조 보정
                h_level = min(6, level + 1)
                full_title = f"{clause_num} {title}".strip() if clause_num else title
                
                sections.append({
                    "clause": clause_num,
                    "title": title,
                    "content": content,
                    "level": level,
                    "page": 1,
                    "headers": {f"H{h_level}": full_title},
                    "header_path": full_title,
                    "pattern": "ai_led_mapping"
                })

                # 서브 섹션들 추가
                for ss in sub_sections:
                    ss_full_title = f"{ss['clause']} {ss['title']}"
                    sections.append({
                        "clause": ss["clause"],
                        "title": ss["title"],
                        "content": ss["content"],
                        "level": ss["level"],
                        "page": 1,
                        "headers": {f"H{h_level}": ss_full_title},
                        "header_path": ss_full_title,
                        "pattern": "ai_led_sub_split"
                    })
            
            if sections:
                # 계층 구조 보완 (부모 찾기)
                for i, sec in enumerate(sections):
                    if i > 0:
                        # 자신보다 레벨이 한 단계 낮은 가장 가까운 이전 섹션을 부모로 지정
                        for j in range(i-1, -1, -1):
                            if sections[j]["level"] < sec["level"]:
                                sec["parent"] = sections[j]["clause"]
                                break
                                
                state["sections"] = sections
                print(f"   ✅ [Split] AI 지도를 기반으로 {len(sections)}개의 조항을 정밀하게 분할했습니다.")
                return state

    # [Fallback] AI 실패 시 헤더(#) 기반 분할
    print("   📋 [Fallback] 헤더(#) 기호 기반으로 분할합니다.")
    # ... (기존 로직 유지)

    # 3. 헤더 기반 파싱 (폴백)
    lines = effective_markdown.split('\n')
    sections = []
    current_headers = {1: None, 2: None, 3: None, 4: None, 5: None, 6: None}
    current_content = []
    current_page = 1
    
    def flush_section():
        nonlocal current_content
        if current_content:
            content = '\n'.join(current_content).strip()
            if content:
                header_path_parts = []
                headers_dict = {}
                for level in range(1, 7):
                    if current_headers[level]:
                        headers_dict[f"H{level}"] = current_headers[level]
                        header_path_parts.append(current_headers[level])
                
                parent = None
                for level in range(6, 0, -1):
                    if current_headers[level]:
                        for p_level in range(level - 1, 0, -1):
                            if current_headers[p_level]:
                                parent = current_headers[p_level]
                                break
                        break
                
                sections.append({
                    "content": content,
                    "headers": headers_dict,
                    "header_path": " > ".join(header_path_parts) if header_path_parts else None,
                    "page": current_page,
                    "parent": parent,
                })
        current_content = []
    
    for line in lines:
        page_match = re.match(r'<!-- PAGE:(\d+) -->', line)
        if page_match:
            current_page = int(page_match.group(1))
            continue
        
        header_match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if header_match:
            flush_section()
            level = len(header_match.group(1))
            header_text = header_match.group(2).strip()
            current_headers[level] = header_text
            for l in range(level + 1, 7):
                current_headers[l] = None
            current_content.append(line)
        else:
            current_content.append(line)
    
    flush_section()
    state["sections"] = sections
    return state


def node_optimize(state: PipelineState) -> PipelineState:
    """
    5단계: 긴 섹션 재분할 + 상세 메타데이터 추출 + 최적화
    """
    sections = state.get("sections", [])
    chunk_size = state.get("chunk_size", 500)
    chunk_overlap = state.get("chunk_overlap", 50)
    doc_meta = state.get("metadata", {})
    use_llm_metadata = state.get("use_llm_metadata", False)
    
    chunks = []
    idx = 0
    
    # 문서 레벨 메타데이터
    doc_id = doc_meta.get("doc_id")
    doc_title = doc_meta.get("title")
    version = doc_meta.get("version")
    
    # 1단계: 메타데이터 추출 및 조항 분석 (안정성을 위해 다시 순차 처리)
    for section in sections:
        content = section["content"]
        headers = section.get("headers", {})
        clause_level = 0
        for l in range(6, 0, -1):
            if headers.get(f"H{l}"):
                clause_level = l
                break
        
        # 조항 번호 및 제목 추출 (파이프라인에서 이미 추출된 것이 있다면 우선 사용)
        current_section_title = section.get("title") or headers.get(f"H{clause_level}") or "Untitled"
        clause_id = section.get("clause")
        
        if not clause_id:
            num_match = re.search(r'([제]?\d+(?:\.\d+)*[조]?)', current_section_title)
            if num_match:
                clause_id = num_match.group(1)
            
        # 상위 섹션 번호 유추 (5.1.2 -> 5)
        main_section = clause_id.split('.')[0] if clause_id and '.' in clause_id else clause_id
        
        # 섹션 정보 저장
        section["main_section"] = main_section
        section["clause_id"] = clause_id
        section["current_title"] = current_section_title
        section["clause_level"] = clause_level

        # 조항별 상세 메타데이터 추출 (AI) - 순차 처리로 복구
        # 제목이 있고, 조항 번호가 있으며, 본문이 100자 이상인 경우에만 분석
        clause_meta = {}
        section_idx = sections.index(section) + 1
        
        if use_llm_metadata and current_section_title != "Untitled" and clause_id and len(content.strip()) > 100:
            print(f"   🔎 [{section_idx}/{len(sections)}] 조항 분석 중: {current_section_title}")
            clause_meta = extract_clause_metadata(content, doc_meta, current_section_title)
        else:
            # AI 분석 스킵하는 경우도 로그에 표시 (짧은 조항)
            if clause_id:
                print(f"   📋 [{section_idx}/{len(sections)}] 조항 저장: {clause_id} {current_section_title[:30] if current_section_title else ''}...")

        section["clause_meta"] = clause_meta
        idx += 1 # 다음 섹션을 위해 인덱스 증가 (누락 방지)

    # 2단계: 최적화 및 청크 생성
    for section in sections:
        content = section["content"]
        clause_id = section.get("clause_id")
        current_section_title = section.get("current_title")
        clause_meta = section.get("clause_meta", {})
        main_section = section.get("main_section")
        clause_level = section.get("clause_level", 0)
        
        # 긴 섹션 재분할
        if len(content) > chunk_size:
            text_chunks = _split_recursive(content, chunk_size, chunk_overlap)
            is_split = len(text_chunks) > 1
        else:
            text_chunks = [content]
            is_split = False
        
        for i, text in enumerate(text_chunks):
            if not text.strip(): continue
            
            section_id = f"{doc_id}:{clause_id}" if clause_id else f"{doc_id}:CH{idx}"
            
            # 고도화된 메타데이터 구조 (V22.0)
            meta = {
                "doc_id": doc_id,
                "doc_title": doc_title,
                "clause_id": clause_id,
                "title": current_section_title,
                "clause_level": clause_level,
                "main_section": main_section,
                "section_id": section_id,
                # LLM 분석 데이터
                "content_type": clause_meta.get("content_type"),
                "main_topic": clause_meta.get("main_topic"),
                "sub_topics": clause_meta.get("sub_topics"),
                "actors": clause_meta.get("actors"),
                "actions": clause_meta.get("actions"),
                "conditions": clause_meta.get("conditions"),
                "summary": clause_meta.get("summary"),
                "intent_scope": clause_meta.get("intent_scope"),
                "intent_summary": clause_meta.get("intent_summary"),
                "language": clause_meta.get("language", "ko"),
                # 시스템 관리 정보
                "page": section.get("page", 1),
                "parent": section.get("parent"),
                "chunk_part": i + 1 if is_split else None,
            }
            
            chunks.append({
                "text": text.strip(),
                "index": idx,
                "metadata": meta
            })
            idx += 1
    
    state["chunks"] = chunks
    return state


def node_finalize(state: PipelineState) -> PipelineState:
    """
    6단계: 최종 결과 정리
    """
    chunks = state.get("chunks", [])
    
    if not chunks:
        state["success"] = False
        state["errors"] = ["청크 생성 실패: 결과가 비어있습니다."]
    else:
        state["success"] = True
    
    # 통계 추가
    state["metadata"]["total_chunks"] = len(chunks)
    state["metadata"]["quality_score"] = state.get("quality_score", 0)
    state["metadata"]["conversion_method"] = state.get("conversion_method", "unknown")
    
    return state


# ═══════════════════════════════════════════════════════════════════════════
# 조건부 라우팅 함수
# ═══════════════════════════════════════════════════════════════════════════

def should_fallback(state: PipelineState) -> Literal["fallback", "validate"]:
    """변환 실패 시 폴백으로 라우팅"""
    if state.get("errors") or not state.get("markdown"):
        return "fallback"
    return "validate"


def should_repair(state: PipelineState) -> Literal["repair", "split"]:
    """품질 점수가 낮으면 보정으로 라우팅"""
    quality_score = state.get("quality_score", 0)
    retry_count = state.get("retry_count", 0)
    
    if quality_score < 0.5 and retry_count < 2:
        return "repair"
    return "split"


# ═══════════════════════════════════════════════════════════════════════════
# 헬퍼 함수들
# ═══════════════════════════════════════════════════════════════════════════

def _convert_docx(filename: str, content: bytes) -> tuple:
    """DOCX → Markdown"""
    from docx import Document
    
    doc = Document(BytesIO(content))
    md_lines = []
    metadata = {}
    
    sop_pattern = re.compile(r'((?:EQ-)?SOP[-_]?\d{4,5})', re.IGNORECASE)
    
    main_sections = ['목적', 'Purpose', '적용 범위', 'Scope', '정의', 'Definitions',
                     '책임', 'Responsibilities', '절차', 'Procedure', 
                     '참고문헌', 'Reference', '첨부', 'Attachments']
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            md_lines.append("")
            continue
        
        # SOP ID 추출
        sop_match = sop_pattern.search(text)
        if sop_match and "sop_id" not in metadata:
            sop_id = sop_match.group(1).upper().replace('_', '-')
            if not sop_id.startswith('EQ-'):
                sop_id = 'EQ-' + sop_id
            metadata["sop_id"] = sop_id
        
        # 헤더 레벨 결정
        header_level = None
        
        style_name = para.style.name.lower() if para.style else ""
        if 'heading 1' in style_name or 'title' in style_name:
            header_level = 1
        elif 'heading 2' in style_name:
            header_level = 2
        elif 'heading 3' in style_name:
            header_level = 3
        elif 'heading 4' in style_name:
            header_level = 4
        
        if not header_level:
            for section in main_sections:
                if text.startswith(section) or re.match(rf'^\d+\s+{section}', text):
                    header_level = 2
                    break
            
            if not header_level:
                if re.match(r'^\d+\.\d+\.\d+\.\d+\s*', text) and len(text) < 60:
                    header_level = 5
                elif re.match(r'^\d+\.\d+\.\d+\s+', text) and len(text) < 60:
                    header_level = 4
                elif re.match(r'^\d+\.\d+\s+', text) and len(text) < 60:
                    header_level = 3
                elif re.match(r'^\d+\.?\s+[가-힣A-Za-z]', text) and len(text) < 60:
                    header_level = 2
                elif re.match(r'^[가-힣A-Z][가-힣\s\(\)/·\-]+\s*\([A-Za-z\s&/\-:]+\)\s*$', text):
                    header_level = 3
        
        if header_level:
            md_lines.append(f"{'#' * header_level} {text}")
        else:
            md_lines.append(text)
    
    # 테이블 처리
    for table in doc.tables:
        md_lines.append("")
        md_lines.append(_table_to_markdown(table))
    
    return '\n'.join(md_lines), metadata


def _convert_pdf_with_fallback(filename: str, content: bytes) -> tuple:
    """PDF 변환 (다중 폴백) + 페이지 마커"""
    
    # 1순위: pdfplumber (가장 안정적)
    try:
        import pdfplumber
        md_lines = []
        total_text_len = 0
        with pdfplumber.open(BytesIO(content)) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text() or ''
                if text.strip():
                    md_lines.append(f"<!-- PAGE:{i + 1} -->")
                    md_lines.append(text)
                    total_text_len += len(text.strip())
        
        # 실제 텍스트가 의미 있는 수준(예: 50자) 이상일 때만 성공으로 간주
        if total_text_len > 50:
            return '\n'.join(md_lines), {"parser": "pdfplumber", "total_pages": len(pdf.pages)}, "pdfplumber"
        print(f"   ⚠️ pdfplumber: 텍스트 추출 부족 ({total_text_len}자). 다른 파서 시도...")
    except Exception as e:
        print(f"   pdfplumber 실패: {e}")
    
    # 2순위: Docling
    try:
        from docling.document_converter import DocumentConverter
        import tempfile
        import os
        
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as f:
            f.write(content)
            temp_path = f.name
        
        try:
            converter = DocumentConverter()
            result = converter.convert(temp_path)
            markdown = result.document.export_to_markdown()
            if len(markdown.strip()) > 50:
                return markdown, {"parser": "docling"}, "docling"
            print(f"   ⚠️ Docling: 텍스트 추출 부족. 다른 파서 시도...")
        finally:
            os.unlink(temp_path)
    except Exception as e:
        print(f"   Docling 실패: {e}")
    
    # 3순위: PyMuPDF
    try:
        import fitz
        pdf = fitz.open(stream=content, filetype="pdf")
        md_lines = []
        total_text_len = 0
        for page_num, page in enumerate(pdf):
            text = page.get_text()
            if text.strip():
                md_lines.append(f"<!-- PAGE:{page_num + 1} -->")
                md_lines.append(text)
                total_text_len += len(text.strip())
        if total_text_len > 50:
            return '\n'.join(md_lines), {"parser": "pymupdf", "total_pages": len(pdf)}, "pymupdf"
        print(f"   ⚠️ PyMuPDF: 텍스트 추출 부족 ({total_text_len}자). 다른 파서 시도...")
    except Exception as e:
        print(f"   PyMuPDF 실패: {e}")
    
    # 4순위: PyPDF2
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(BytesIO(content))
        md_lines = []
        total_text_len = 0
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ''
            if text.strip():
                md_lines.append(f"<!-- PAGE:{i + 1} -->")
                md_lines.append(text)
                total_text_len += len(text.strip())
        if total_text_len > 50:
            return '\n'.join(md_lines), {"parser": "pypdf2", "total_pages": len(reader.pages)}, "pypdf2"
        print(f"   ⚠️ PyPDF2: 텍스트 추출 부족 ({total_text_len}자). 다른 파서 시도...")
    except Exception as e:
        print(f"   PyPDF2 실패: {e}")
    
    # 5순위: OCR Fallback (스캔본/이미지 PDF용)
    try:
        print("   스캔 문서 감지: OCR(광학 문자 인식) 엔진 가동 중...")
        markdown, metadata = _convert_pdf_ocr(content)
        if len(markdown.strip()) > 50:
            return markdown, {**metadata, "parser": "ocr"}, "ocr"
    except Exception as e:
        print(f"   OCR 파서 실패: {e}")
    
    raise Exception("모든 PDF 파서 실패 (OCR 포함)")

def _convert_pdf_ocr(content: bytes) -> tuple:
    """PDF OCR 처리 (Tesseract 기반)"""
    try:
        import pytesseract
        from pdf2image import convert_from_bytes
        from PIL import Image
        
        # PDF를 이미지로 변환 (300 DPI 권장)
        pages = convert_from_bytes(content, dpi=300)
        
        md_lines = []
        for i, page in enumerate(pages):
            # 한글 + 영어 OCR 수행
            text = pytesseract.image_to_string(page, lang='kor+eng')
            if text.strip():
                md_lines.append(f"<!-- PAGE:{i + 1} (OCR) -->")
                md_lines.append(text)
        
        return '\n'.join(md_lines), {"total_pages": len(pages)}
    except ImportError:
        return "OCR 필요한 라이브러리(pytesseract, pdf2image)가 설치되지 않았습니다.", {}
    except Exception as e:
        return f"OCR 처리 중 오류 발생: {str(e)}", {}


def _convert_html(filename: str, content: bytes) -> tuple:
    """HTML → Markdown"""
    from bs4 import BeautifulSoup
    
    html = content.decode('utf-8', errors='ignore')
    soup = BeautifulSoup(html, 'html.parser')
    
    for tag in soup(['script', 'style', 'nav', 'footer', 'header']):
        tag.decompose()
    
    md_lines = []
    title = soup.title.string if soup.title else filename
    md_lines.append(f"# {title}")
    md_lines.append("")
    
    for tag in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'li']):
        if tag.name.startswith('h'):
            level = int(tag.name[1])
            md_lines.append(f"{'#' * level} {tag.get_text(strip=True)}")
        elif tag.name == 'li':
            md_lines.append(f"- {tag.get_text(strip=True)}")
        else:
            text = tag.get_text(strip=True)
            if text:
                md_lines.append(text)
        md_lines.append("")
    
    return '\n'.join(md_lines), {"title": title}


def _convert_text_to_markdown(text: str) -> str:
    """텍스트 → 마크다운 (헤더 추론)"""
    return _infer_headers(text)


def _pdf_fallback_extract(content: bytes) -> str:
    """PDF 폴백 추출"""
    try:
        import pdfplumber
        with pdfplumber.open(BytesIO(content)) as pdf:
            texts = []
            for i, page in enumerate(pdf.pages):
                texts.append(f"<!-- PAGE:{i+1} -->")
                texts.append(page.extract_text() or '')
            return '\n\n'.join(texts)
    except:
        pass
    return content.decode('latin-1', errors='ignore')


def _docx_fallback_extract(content: bytes) -> str:
    """DOCX 폴백: XML 직접 파싱"""
    import zipfile
    from xml.etree import ElementTree
    
    try:
        with zipfile.ZipFile(BytesIO(content)) as zf:
            xml_content = zf.read('word/document.xml')
            tree = ElementTree.fromstring(xml_content)
            
            texts = []
            for t in tree.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                if t.text:
                    texts.append(t.text)
            
            return '\n'.join(texts)
    except:
        return ""


def _clean_noise_globally(markdown: str) -> str:
    """
    전체 문서에서 반복되는 SOP 헤더, 푸터, 메타데이터 블록을 제거합니다.
    (V22.12 Global Filter - 문서 끝부분 제거 추가)
    """
    if not markdown:
        return ""
    
    # [V22.12] 문서 끝부분 감지 및 제거 (개정이력, 승인정보)
    end_markers = [
        '문서개정이력',
        'Document Revision History',
        'Document Approvals',
        '***END OF DOCUMENT***',
    ]
    
    # 문서 끝부분 시작 위치 찾기
    end_pos = len(markdown)
    for marker in end_markers:
        idx = markdown.find(marker)
        if idx > 0 and idx < end_pos:
            end_pos = idx
    
    # 문서 끝부분 제거 (본문만 유지)
    if end_pos < len(markdown):
        markdown = markdown[:end_pos].strip()
        
    lines = markdown.split('\n')
    cleaned_lines = []
    
    # 제거할 패턴들 (대소문자 무시)
    noise_patterns = [
        r'^Number:\s*EQ-SOP-\d+.*',
        r'^Version:\s*\d+\.\d+.*',
        r'^Effective Date:\s*\d{4}-\d{2}-\d{2}.*',
        r'^Owning Department\s*.*',
        r'^Title\s*.*',
        r'^\d+\s+of\s+\d+$', # 페이지 번호 (예: 1 of 11)
        r'^\*\*\*END OF DOCUMENT\*\*\*$',
    ]
    
    # 문서 제목 중복 제거를 위한 추적 (첫 3줄 정도에서 제목 추출 시도)
    doc_title = None
    for i in range(min(10, len(lines))):
        line = lines[i].strip()
        if line and not any(re.match(p, line, re.I) for p in noise_patterns) and not line.startswith('#') and not line.startswith('<!--'):
            doc_title = line
            break
            
    for line in lines:
        stripped = line.strip()
        if not stripped:
            cleaned_lines.append(line)
            continue
            
        # 1. 명시적 노이즈 패턴 체크
        is_noise = False
        for pattern in noise_patterns:
            if re.match(pattern, stripped, re.I):
                is_noise = True
                break
        
        if is_noise:
            continue
            
        # 2. 반복되는 문서 제목 체크 (메타데이터 근처에 있는 경우)
        if doc_title and stripped == doc_title:
            # 첫 번째 제목은 유지, 나머지는 제거 (단, 헤더 # 가 붙어있지 않은 경우만)
            # (보통 상단 표 안에 제목이 반복되므로 이를 제거하는 목적)
            if cleaned_lines.count(line) > 0 and not line.startswith('#'):
                continue
                
        cleaned_lines.append(line)
        
    return '\n'.join(cleaned_lines)


def _table_to_markdown(table) -> str:
    """Word 테이블 → Markdown"""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
        rows.append(cells)
    
    if not rows:
        return ""
    
    md_lines = []
    md_lines.append("| " + " | ".join(rows[0]) + " |")
    md_lines.append("| " + " | ".join(["---"] * len(rows[0])) + " |")
    for row in rows[1:]:
        while len(row) < len(rows[0]):
            row.append("")
        md_lines.append("| " + " | ".join(row[:len(rows[0])]) + " |")
    
    return '\n'.join(md_lines)


def _infer_headers(markdown: str) -> str:
    """헤더 추론 삽입 (PDF용 강화) - V22.11"""
    lines = markdown.split('\n')
    result = []
    
    main_sections = ['목적', '적용 범위', '정의', '책임', '절차', '참고문헌', '첨부',
                     'Purpose', 'Scope', 'Definitions', 'Responsibilities', 'Procedure', 
                     'Reference', 'Attachments']
    
    # 무시할 패턴 (페이지 번호 등) - 텍스트는 유지하되 헤더로 안 만듦
    ignore_patterns = [
        r'^\d+\s+of\s+\d+$',
        r'^Page\s+\d+',
        r'^-\s*\d+\s*-$',
        r'^Number:\s*',
        r'^<!--\s*PAGE',
    ]
    
    for line in lines:
        stripped = line.strip()
        
        if not stripped:
            result.append(line)
            continue
        
        # 이미 마크다운 헤더가 붙어있으면 스킵
        if stripped.startswith('#'):
            result.append(line)
            continue
        
        # 무시 패턴 체크 (헤더로 안 만들고 텍스트 유지)
        should_ignore = False
        for pattern in ignore_patterns:
            if re.match(pattern, stripped, re.IGNORECASE):
                should_ignore = True
                break
        
        if should_ignore:
            result.append(line)
            continue
        
        # 1. 숫자형 헤더 패턴 (V22.11: 공백 또는 탭 허용, 더 유연하게)
        # 5.1.2.1.1 xxx → H6
        if re.match(r'^(\d+\.\d+\.\d+\.\d+\.\d+)\s*(.+)', stripped):
            result.append(f"###### {stripped}")
            continue
        
        # 5.1.2.1 xxx → H5
        if re.match(r'^(\d+\.\d+\.\d+\.\d+)\s*(.+)', stripped):
            result.append(f"##### {stripped}")
            continue
        
        # 5.1.1 xxx → H4
        if re.match(r'^(\d+\.\d+\.\d+)\s*(.+)', stripped):
            result.append(f"#### {stripped}")
            continue
        
        # 5.1 xxx → H3 (조항 번호 뒤에 바로 한글/영문이 오는 경우도 허용)
        if re.match(r'^(\d+\.\d+)\s*([가-힣A-Za-z].+)', stripped):
            result.append(f"### {stripped}")
            continue
        
        # 5 xxx → H2 (단, "X of Y" 같은 페이지 번호 제외)
        match = re.match(r'^(\d+)\s+([가-힣A-Za-z].+)', stripped)
        if match:
            num = match.group(1)
            text = match.group(2)
            if not re.match(r'^of\s+\d+', text, re.IGNORECASE):
                result.append(f"## {stripped}")
                continue
        
        # 2. 주요 섹션 키워드 → H2 (V22.11: 숫자로 시작하는 경우에만 허용)
        # "목적으로 한다" 같은 일반 문장 방어
        matched = False
        for section in main_sections:
            # [핵심] 숫자+공백+키워드 형태만 헤더로 인정 (예: "1 목적")
            # 순수 키워드만 있는 경우는 이미 위에서 "5 xxx" 패턴으로 처리됨
            pattern = rf'^(\d+)\s+{re.escape(section)}(\s+|:|$)'
            if re.match(pattern, stripped) and len(stripped) < 50:
                result.append(f"## {stripped}")
                matched = True
                break
        
        if not matched:
            # 3. 소제목 패턴 → H3 (한글 제목 + 영문 괄호)
            if re.match(r'^[가-힣][가-힣\s\(\)/·\-]+\s*\([A-Za-z\s&/\-:]+\)\s*$', stripped):
                result.append(f"### {stripped}")
            else:
                result.append(line)
    
    return '\n'.join(result)


def _repair_tables(markdown: str) -> str:
    """깨진 테이블 복구"""
    lines = markdown.split('\n')
    result = []
    in_table = False
    table_cols = 0
    
    for line in lines:
        if line.strip().startswith('|') and line.strip().endswith('|'):
            cols = line.count('|') - 1
            
            if not in_table:
                in_table = True
                table_cols = cols
                result.append(line)
            else:
                while line.count('|') - 1 < table_cols:
                    line = line.rstrip('|') + ' |'
                result.append(line)
        else:
            in_table = False
            result.append(line)
    
    return '\n'.join(result)


def _split_recursive(text: str, chunk_size: int, overlap: int) -> List[str]:
    """재귀적 텍스트 분할"""
    if len(text) <= chunk_size:
        return [text]
    
    separators = ["\n\n", "\n| ", "\n", ". ", "。", " ", ""]
    is_table = text.strip().startswith('|') or '\n|' in text
    effective_overlap = 0 if is_table else overlap
    
    for sep in separators:
        if sep in text:
            parts = text.split(sep)
            chunks = []
            current = ""
            
            for part in parts:
                if len(current) + len(part) + len(sep) <= chunk_size:
                    current = current + sep + part if current else part
                else:
                    if current:
                        chunks.append(current)
                    if len(part) > chunk_size:
                        chunks.extend(_split_recursive(part, chunk_size, overlap))
                        current = ""
                    else:
                        current = part
            
            if current:
                chunks.append(current)
            
            return chunks
    
    step = chunk_size - effective_overlap if effective_overlap > 0 else chunk_size
    return [text[i:i+chunk_size] for i in range(0, len(text), step)]


# ═══════════════════════════════════════════════════════════════════════════
# LangGraph 파이프라인 빌더
# ═══════════════════════════════════════════════════════════════════════════

def build_pipeline():
    """LangGraph 파이프라인 구성"""
    try:
        from langgraph.graph import StateGraph, END
    except ImportError:
        raise ImportError("langgraph 패키지가 필요합니다: pip install langgraph")
    
    workflow = StateGraph(PipelineState)
    
    workflow.add_node("load", node_load)
    workflow.add_node("convert", node_convert)
    workflow.add_node("fallback", node_convert_fallback)
    workflow.add_node("validate", node_validate)
    workflow.add_node("repair", node_repair)
    workflow.add_node("split", node_split)
    workflow.add_node("optimize", node_optimize)
    workflow.add_node("finalize", node_finalize)
    
    workflow.set_entry_point("load")
    
    workflow.add_edge("load", "convert")
    
    workflow.add_conditional_edges(
        "convert",
        should_fallback,
        {"fallback": "fallback", "validate": "validate"}
    )
    
    workflow.add_edge("fallback", "validate")
    
    workflow.add_conditional_edges(
        "validate",
        should_repair,
        {"repair": "repair", "split": "split"}
    )
    
    workflow.add_edge("repair", "split")
    workflow.add_edge("split", "optimize")
    workflow.add_edge("optimize", "finalize")
    workflow.add_edge("finalize", END)
    
    return workflow.compile()


# ═══════════════════════════════════════════════════════════════════════════
# 메인 함수
# ═══════════════════════════════════════════════════════════════════════════

def process_document(
    filename: str,
    content: bytes,
    chunk_size: int = 500,
    chunk_overlap: int = 50,
    debug: bool = False,
    use_llm_metadata: bool = False,  # LLM 메타데이터 추출 사용 여부
    use_clause_parsing: bool = True  # 조항 번호 기반 파싱 사용 여부
) -> dict:
    """문서 처리 메인 함수"""
    
    initial_state: PipelineState = {
        "filename": filename,
        "content": content,
        "chunk_size": chunk_size,
        "chunk_overlap": chunk_overlap,
        "use_llm_metadata": use_llm_metadata,  # LLM 메타데이터 추출 옵션
        "use_clause_parsing": use_clause_parsing,  # 조항 번호 기반 파싱 옵션
        "file_type": "",
        "markdown": "",
        "metadata": {},
        "sections": [],
        "chunks": [],
        "quality_score": 0.0,
        "conversion_method": "",
        "errors": [],
        "warnings": [],
        "retry_count": 0,
        "success": False,
    }
    
    try:
        pipeline = build_pipeline()
        result = pipeline.invoke(initial_state)
        
        if debug:
            print(f"\n{'='*60}")
            print(f"LangGraph 파이프라인 결과")
            print(f"{'='*60}")
            print(f"   파일: {filename}")
            print(f"   변환 방법: {result.get('conversion_method')}")
            print(f"   품질 점수: {result.get('quality_score', 0):.0%}")
            print(f"   총 청크: {len(result.get('chunks', []))}")
            if result.get('warnings'):
                print(f"   ⚠️ 경고: {result['warnings']}")
            if result.get('errors'):
                print(f"   ❌ 에러: {result['errors']}")
        
        return result
        
    except ImportError:
        if debug:
            print("⚠️ LangGraph 없음, 심플 파이프라인 사용")
        return _simple_pipeline(initial_state, debug)


def _simple_pipeline(state: PipelineState, debug: bool = False) -> dict:
    """LangGraph 없을 때 사용하는 심플 파이프라인"""
    state = node_load(state)
    if state.get("errors"):
        return state
    
    state = node_convert(state)
    if not state.get("markdown"):
        state = node_convert_fallback(state)
    
    state = node_validate(state)
    
    if state.get("quality_score", 0) < 0.5:
        state = node_repair(state)
    
    state = node_split(state)
    state = node_optimize(state)
    state = node_finalize(state)
    
    if debug:
        print(f"\n심플 파이프라인 결과: {len(state.get('chunks', []))} 청크")
    
    return state


# ═══════════════════════════════════════════════════════════════════════════
# 결과 변환 헬퍼
# ═══════════════════════════════════════════════════════════════════════════

def state_to_chunks(state: dict) -> List[Chunk]:
    """상태를 Chunk 객체 리스트로 변환"""
    return [
        Chunk(
            text=c["text"],
            index=c["index"],
            metadata=c["metadata"]
        )
        for c in state.get("chunks", [])
    ]


# ═══════════════════════════════════════════════════════════════════════════
# 테스트
# ═══════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    print("document_pipeline v9.1 테스트")