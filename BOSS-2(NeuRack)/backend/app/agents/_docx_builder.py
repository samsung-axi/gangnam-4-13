"""Fills admin application DOCX templates with user-supplied field values."""

import io
import re
from pathlib import Path

from docx import Document

_ASSETS = Path(__file__).parent.parent / "assets"
_BIZ_REG_TEMPLATE = _ASSETS / "business_registration_template.docx"
_MAIL_ORDER_TEMPLATE = _ASSETS / "mail_order_template.docx"
_PURCHASE_SAFETY_TEMPLATE = _ASSETS / "purchase_safety_template.docx"

# ── 사업자등록 신청서 ─────────────────────────────────────────────────────────
# (row_idx, col_idx) → field key  (Table 0)
_BIZ_REG_CELL_MAP: dict[tuple[int, int], str] = {
    (6, 1): "business_name",
    (7, 1): "representative_name",
    (6, 14): "phone_business",
    (8, 21): "phone_mobile",
    (10, 3): "location",
    (14, 4): "industry_type",
    (14, 9): "industry_item",
    (15, 21): "industry_code",
    (16, 26): "opening_date",
    (16, 29): "employees_count",
    (18, 1): "cyber_mall_name",
    (18, 18): "cyber_mall_domain",
    (21, 1): "owned_area",
    (21, 5): "rented_area",
    (21, 7): "landlord_name",
    (21, 9): "landlord_reg_no",
    (21, 18): "lease_period",
    (21, 22): "lease_deposit",
    (21, 28): "lease_monthly",
    (25, 3): "own_capital",
    (25, 18): "borrowed_capital",
    (27, 1): "email",
    (30, 14): "신청일",
    (30, 20): "representative_name",
}

# ── 통신판매업 신고서 ─────────────────────────────────────────────────────────
# label cells: value is appended as a new paragraph below the label
_MAIL_ORDER_LABEL_CELLS: list[tuple[int, int, str]] = [
    (5, 2, "business_name"),       # 법인명(상호)
    (6, 2, "location"),            # 소재지
    (6, 9, "phone_business"),      # 전화번호(법인)
    (7, 2, "representative_name"), # 대표자의 성명
    (8, 2, "location"),            # 주소(대표자) — same field
    (8, 9, "phone_mobile"),        # 전화번호(대표자)
    (9, 2, "email"),               # 전자우편주소
    (9, 9, "business_reg_no"),     # 사업자등록번호
    (10, 2, "internet_domain"),    # 인터넷도메인
    (10, 9, "host_server_location"), # 호스트서버 소재지
]


def _set_cell_text(cell, text: str) -> None:
    """Replace paragraph 0 text, preserving run formatting."""
    para = cell.paragraphs[0]
    fmt: dict = {}
    if para.runs:
        r0 = para.runs[0]
        fmt = {
            "bold": r0.bold,
            "font_name": r0.font.name,
            "font_size": r0.font.size,
        }
    for r in para.runs:
        r._element.getparent().remove(r._element)
    run = para.add_run(text)
    if fmt:
        run.bold = fmt.get("bold")
        if fmt.get("font_name"):
            run.font.name = fmt["font_name"]
        if fmt.get("font_size"):
            run.font.size = fmt["font_size"]


def _append_value_para(cell, value: str) -> None:
    """Add a new paragraph with the value below existing label content."""
    if not value:
        return
    para = cell.add_paragraph(value)
    # inherit font from label paragraph if possible
    if len(cell.paragraphs) >= 2:
        label_para = cell.paragraphs[0]
        if label_para.runs:
            r0 = label_para.runs[0]
            run = para.runs[0] if para.runs else para.add_run(value)
            if not para.runs:
                run = para.add_run(value)
            if r0.font.name:
                run.font.name = r0.font.name
            if r0.font.size:
                run.font.size = r0.font.size


def _replace_para_text(para, new_text: str) -> None:
    """Replace all runs in a paragraph with a single new-text run."""
    for r in list(para.runs):
        r._element.getparent().remove(r._element)
    para.add_run(new_text)


def _format_date_korean(date_str: str) -> str:
    """Convert YYYY-MM-DD or YYYY년MM월DD일 → 'YYYY년  MM월  DD일' for form."""
    m = re.match(r"(\d{4})-(\d{2})-(\d{2})", date_str or "")
    if m:
        return f"{m.group(1)}년  {m.group(2)}월  {m.group(3)}일"
    return date_str


# ── builders ──────────────────────────────────────────────────────────────────

def build_docx(fields: dict[str, str], doc_type: str = "business_registration") -> bytes:
    if doc_type == "mail_order_registration":
        return _build_mail_order(fields)
    if doc_type == "purchase_safety_exempt":
        return _build_purchase_safety(fields)
    return _build_biz_reg(fields)


def _build_biz_reg(fields: dict[str, str]) -> bytes:
    doc = Document(str(_BIZ_REG_TEMPLATE))
    tbl = doc.tables[0]
    seen: set[int] = set()
    for (ri, ci), key in _BIZ_REG_CELL_MAP.items():
        val = fields.get(key, "")
        if not val:
            continue
        cell = tbl.rows[ri].cells[ci]
        tc_id = id(cell._tc)
        if tc_id in seen:
            continue
        seen.add(tc_id)
        _set_cell_text(cell, val)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_mail_order(fields: dict[str, str]) -> bytes:
    doc = Document(str(_MAIL_ORDER_TEMPLATE))
    tbl = doc.tables[0]

    seen: set[int] = set()
    for ri, ci, key in _MAIL_ORDER_LABEL_CELLS:
        val = fields.get(key, "")
        if not val:
            continue
        cell = tbl.rows[ri].cells[ci]
        tc_id = id(cell._tc)
        if tc_id in seen:
            continue
        seen.add(tc_id)
        _append_value_para(cell, val)

    # 날짜 (row 14, col 0): "년     월     일" → actual date
    date_val = fields.get("신청일", "")
    if date_val:
        date_cell = tbl.rows[14].cells[0]
        _replace_para_text(date_cell.paragraphs[0], _format_date_korean(date_val))

    # 신고인 서명란 (row 15, col 0): append name after "신고인"
    name_val = fields.get("representative_name", "")
    if name_val:
        sign_cell = tbl.rows[15].cells[0]
        tc_id = id(sign_cell._tc)
        if tc_id not in seen:
            seen.add(tc_id)
            _replace_para_text(sign_cell.paragraphs[0], f"신고인: {name_val}")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_purchase_safety(fields: dict[str, str]) -> bytes:
    doc = Document(str(_PURCHASE_SAFETY_TEMPLATE))
    cell = doc.tables[0].rows[0].cells[0]

    name_val = fields.get("representative_name", "")
    date_val = fields.get("신청일", "")

    for para in cell.paragraphs:
        txt = para.text
        # 신고인 서명 줄: "  신고인                       는 ｢..."
        if name_val and "신고인" in txt and "｢전자상거래" in txt:
            new_txt = re.sub(r"(신고인)\s+(는)", rf"\1 {name_val} \2", txt)
            _replace_para_text(para, new_txt)
        # 날짜 줄: "  년       월       일 "
        elif date_val and re.search(r"년\s+월\s+일", txt):
            _replace_para_text(para, "  " + _format_date_korean(date_val))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
